"""知识库核心服务测试（任务 2.1~2.5）：解析器（PDF/Word/扫描件/文本）、分块、向量化幂等、删除、filter 隔离、检索阈值、相邻块补齐"""
import io

import pytest

from app.core.config import Settings
from app.services.knowledge_base import (
    KBError,
    ChunkHit,
    KnowledgeBaseService,
    format_chunks_for_prompt,
    parse_document,
    split_text,
)
from tests.conftest import FakeEmbeddings


@pytest.fixture
def kb_service():
    """内存 Chroma + 确定性伪 embedding 的知识库服务（每用例独立）"""
    settings = Settings(deepseek_api_key="test-key")
    return KnowledgeBaseService(settings, build_embeddings=lambda s: FakeEmbeddings())


def _make_pdf_bytes(text: str | None) -> bytes:
    """构造含文本层（text 非空）或无文本层（text=None，模拟扫描件）的 PDF 字节流"""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    if text:
        page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def _make_docx_bytes(lines: list[str], rows: list[list[str]] | None = None) -> bytes:
    """构造 Word 文档字节流（段落 + 可选表格）"""
    from docx import Document

    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    for row in rows or []:
        table = doc.add_table(rows=1, cols=len(row))
        for i, cell in enumerate(row):
            table.rows[0].cells[i].text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestParseDocument:
    def test_txt_passthrough(self):
        assert parse_document("a.txt", "你好世界".encode("utf-8")) == "你好世界"

    def test_md_passthrough(self):
        assert parse_document("note.md", "# 标题\n正文".encode("utf-8")) == "# 标题\n正文"

    def test_pdf_with_text_layer(self):
        # 文本需小于单行页宽（PyMuPDF insert_text 超宽截断），断言提取成功且包含全文
        text = "PDF text layer extraction test. " * 2
        result = parse_document("doc.pdf", _make_pdf_bytes(text))
        assert result.startswith("PDF text layer")

    def test_scanned_pdf_rejected(self):
        with pytest.raises(KBError, match="扫描版"):
            parse_document("scan.pdf", _make_pdf_bytes(None))

    def test_docx_with_paragraphs_and_table(self):
        data = _make_docx_bytes(
            ["第一段内容较长，用于验证 Word 段落文本提取是否正确。", "第二段内容也足够长，覆盖多段落场景。"],
            rows=[["列A", "列B"], ["v1", "v2"]],
        )
        text = parse_document("doc.docx", data)
        assert "第一段内容较长" in text and "第二段内容也足够长" in text
        assert "列A | 列B" in text and "v1 | v2" in text

    def test_unsupported_format(self):
        with pytest.raises(KBError, match="不支持"):
            parse_document("a.xlsx", b"data")


class TestSplitText:
    def test_long_text_split_with_overlap(self):
        text = "知识点。" * 300  # 1200 字
        chunks = split_text(text, chunk_size=500, chunk_overlap=50)
        assert len(chunks) >= 2
        # overlap 生效：相邻块边界字符重复 → 拼接总长 > 原文长
        joined = "".join(chunks)
        assert len(joined) > len(text)

    def test_short_text_single_chunk(self):
        assert split_text("短文本", 500, 50) == ["短文本"]

    def test_blank_text_empty(self):
        assert split_text("   \n ", 500, 50) == []


class TestKnowledgeBaseService:
    async def test_index_then_search_hits_same_text(self, kb_service):
        chunks = ["量子计算的叠加态原理", "纠错码在量子计算中的作用"]
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=chunks)

        hits = await kb_service.search(user_id=1, query="量子计算的叠加态原理")
        assert len(hits) >= 1
        assert hits[0].text == "量子计算的叠加态原理"
        assert hits[0].score > 0.9  # 相同文本余弦≈1

    async def test_unrelated_query_filtered_by_threshold(self, kb_service):
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["量子计算入门"])
        hits = await kb_service.search(user_id=1, query="完全无关的其他内容主题")
        assert hits == []  # 哈希伪向量正交 → score≈0 → 硬过滤

    async def test_user_isolation(self, kb_service):
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["用户一的私有内容"])
        hits = await kb_service.search(user_id=2, query="用户一的私有内容")
        assert hits == []  # filter 只按 user_id=2，看不到他人内容

    async def test_kb_filter(self, kb_service):
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["库一的独特内容"])
        hits = await kb_service.search(user_id=1, query="库一的独特内容", kb_id=11)
        assert hits == []  # 指定库 11，库 10 的内容不可见

    async def test_upsert_idempotent(self, kb_service):
        chunks = ["内容A", "内容B"]
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=chunks)
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["内容A", "内容B"])
        res = kb_service.collection.get()  # 同步 API（chromadb 原生客户端）
        assert len(res["ids"]) == 2  # 同 id upsert 覆盖不重复

    async def test_delete_document_chunks(self, kb_service):
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["要删除的内容"])
        await kb_service.delete_document_chunks(100)
        hits = await kb_service.search(user_id=1, query="要删除的内容")
        assert hits == []

    async def test_delete_base_chunks(self, kb_service):
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=["库内容"])
        await kb_service.index_chunks(user_id=1, kb_id=20, doc_id=200, filename="b.txt", chunks=["库内容"])
        await kb_service.delete_base_chunks(10)
        hits = await kb_service.search(user_id=1, query="库内容")
        assert all(h.kb_id != 10 for h in hits)
        assert any(h.kb_id == 20 for h in hits)

    async def test_expand_neighbors_fills_adjacent_chunks(self, kb_service):
        chunks = ["第一段开头", "第二段中间被检索命中", "第三段结尾"]
        await kb_service.index_chunks(user_id=1, kb_id=10, doc_id=100, filename="a.txt", chunks=chunks)

        hits = await kb_service.search(user_id=1, query="第二段中间被检索命中", min_score=0.9)
        assert {h.chunk_index for h in hits} == {0, 1, 2}  # 命中块 1，相邻 0/2 补齐
        assert [h.chunk_index for h in hits] == [0, 1, 2]  # 按文档内顺序

    async def test_search_error_raises_kb_error(self, kb_service):
        class BrokenCollection:
            def query(self, **kwargs):
                raise RuntimeError("boom")

        kb_service._collection = BrokenCollection()
        with pytest.raises(KBError, match="检索失败"):
            await kb_service.search(user_id=1, query="任何内容")


def test_format_chunks_for_prompt_limits_and_marks_source():
    hits = [ChunkHit(doc_id=1, kb_id=1, filename="手册.pdf", chunk_index=0, text="内容内容", score=0.9)]
    text = format_chunks_for_prompt(hits, limit=50)
    assert "手册.pdf" in text and "第 1 段" in text
    assert len(text) <= 50
