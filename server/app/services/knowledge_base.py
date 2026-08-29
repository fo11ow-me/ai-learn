"""私有知识库服务（RAG）：文档解析 → 分块 → 向量化 → Chroma 检索。
MySQL 为权威源（全文），Chroma 为可重建派生索引（写顺序：先 MySQL 后 Chroma，编排在路由任务层）。
依赖注入仿 SearchClient（build_embeddings 可替换）——测试注入 FakeEmbeddings + 内存 Chroma，不依赖真实百炼 key 与网络"""
import asyncio
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import Settings
from app.core.tracing import task_id_kv

_logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "md", "txt"}  # 上传格式白名单（扩展名，小写）
COLLECTION_NAME = "kb_chunks"  # 单 collection + metadata filter 隔离（设计 D2）
MIN_TEXT_CHARS = 50  # 低于该长度视为扫描件/无文本（PDF 无文本层或空文档）


class KBError(Exception):
    """知识库通用错误（解析/向量化/检索失败，携带用户可读原因）"""

    def __init__(self, message: str):
        super().__init__(message)
        self.message = message


@dataclass
class ChunkHit:
    """检索命中片段（含分数与来源元数据，供判定与出题拼接）"""

    doc_id: int
    kb_id: int
    filename: str
    chunk_index: int
    text: str
    score: float


def parse_document(filename: str, data: bytes) -> str:
    """按扩展名解析文档为全文；扫描版/无文本层抛 KBError。OCR 不做，明确拒绝并提示用户转可复制文本。"""
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf":
        return _parse_pdf(data)
    if ext == "docx":
        return _parse_docx(data)
    if ext in ("md", "txt"):
        return data.decode("utf-8", errors="replace")
    raise KBError(f"不支持的文件格式：{ext}")


def _parse_pdf(data: bytes) -> str:
    """PyMuPDF 提取文本层；无文本层（扫描件）拒绝"""
    import fitz  # PyMuPDF（延迟导入：无该依赖时仅上传功能受影响）

    doc = fitz.open(stream=data, filetype="pdf")
    text = "\n".join(page.get_text() for page in doc).strip()
    if len(text) < MIN_TEXT_CHARS:
        raise KBError("该 PDF 为扫描版或无文本层，无法提取文字，请上传可复制文本的 PDF 或 Word 文档")
    return text


def _parse_docx(data: bytes) -> str:
    """python-docx 提取段落与表格（表格转「列1 | 列2」行，保住表格语义）"""
    import io

    from docx import Document  # python-docx（延迟导入）

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n".join(line for line in lines if line.strip()).strip()
    if len(text) < MIN_TEXT_CHARS:
        raise KBError("该 Word 文档无有效文本内容，请检查后重试")
    return text


def split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list[str]:
    """分块。RecursiveCharacterTextSplitter 按段落/标点优先切分，中文按字符计数。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len
    )
    return [c for c in splitter.split_text(text) if c.strip()]


def _default_build_embeddings(settings: Settings):
    """默认 Embedding 构建（百炼 qwen3.7-text-embedding，OpenAI 兼容端点，维度 1024）。
    check_embedding_ctx_length=False：官方 SDK 对未知模型名默认按长文本模型检查上下文，
    对非 OpenAI 模型会误报超长拒绝；chunk_size=批量上限 20。"""
    from langchain_openai import OpenAIEmbeddings

    return OpenAIEmbeddings(
        model=settings.embedding_model,
        api_key=settings.embedding_api_key,
        base_url=settings.embedding_base_url,
        check_embedding_ctx_length=False,
        chunk_size=settings.embedding_batch_size,
    )


class KnowledgeBaseService:
    """Chroma 向量索引与检索（设计 D2：单 collection + metadata filter 隔离）。

    构建参数可注入：测试注入 FakeEmbeddings + 内存 Chroma，生产走 PersistentClient 目录。
    """

    def __init__(
        self,
        settings: Settings,
        build_embeddings: Callable[[Settings], Any] | None = None,
        chroma_dir: str | None = None,
    ):
        self._settings = settings
        self._build_embeddings = build_embeddings or _default_build_embeddings
        self._chroma_dir = chroma_dir if chroma_dir is not None else (settings.kb_chroma_dir or None)
        self._embeddings = None
        self._collection = None

    @property
    def embeddings(self):
        if self._embeddings is None:
            self._embeddings = self._build_embeddings(self._settings)
        return self._embeddings

    @property
    def collection(self):
        """collection 惰性初始化：首次使用时创建；cosine 空间使 relevance = 1 - distance。"""
        if self._collection is None:
            import chromadb

            if self._chroma_dir:
                client = chromadb.PersistentClient(path=self._chroma_dir)
            else:
                client = chromadb.Client()  # 内存模式（测试用；生产需配置 KB_CHROMA_DIR 持久化）
            self._collection = client.get_or_create_collection(
                name=COLLECTION_NAME, metadata={"hnsw:space": "cosine"}
            )
        return self._collection

    async def index_chunks(self, user_id: int, kb_id: int, doc_id: int, filename: str, chunks: list[str]) -> int:
        """向量化并写入 Chroma（幂等 upsert，确定性 id doc_{doc_id}_{idx}），返回写入条数"""
        if not chunks:
            return 0
        ids = [f"doc_{doc_id}_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "user_id": user_id,
                "kb_id": kb_id,
                "doc_id": doc_id,
                "chunk_index": i,
                "filename": filename,
                "source": "kb",
            }
            for i in range(len(chunks))
        ]
        try:
            vectors = await asyncio.to_thread(self.embeddings.embed_documents, chunks)
            await asyncio.to_thread(
                self.collection.upsert, ids=ids, documents=chunks, metadatas=metadatas, embeddings=vectors
            )
        except Exception as exc:
            _logger.warning("%skb index fail doc_id=%s error=%s", task_id_kv(), doc_id, exc, exc_info=exc)
            raise KBError(f"知识库向量化失败：{exc}") from exc
        _logger.info("%skb index ok doc_id=%s chunks=%d", task_id_kv(), doc_id, len(chunks))
        return len(chunks)

    async def delete_document_chunks(self, doc_id: int) -> None:
        """按 doc_id 删除全部 chunk（覆盖更新/删除文档时调用；幂等）"""
        await asyncio.to_thread(self.collection.delete, where={"doc_id": doc_id})
        _logger.info("%skb delete doc_id=%s", task_id_kv(), doc_id)

    async def delete_base_chunks(self, kb_id: int) -> None:
        """按 kb_id 删除知识库全部 chunk（删除库时调用；幂等）"""
        await asyncio.to_thread(self.collection.delete, where={"kb_id": kb_id})
        _logger.info("%skb delete kb_id=%s", task_id_kv(), kb_id)

    async def search(
        self,
        user_id: int,
        query: str,
        kb_id: int | None = None,
        top_k: int | None = None,
        min_score: float | None = None,
    ) -> list[ChunkHit]:
        """语义检索（设计 D2 骨架）：filter 隔离（user_id 必选，kb_id 可选）→ 余弦距离转 relevance
        → 硬过滤阈值 → 相邻块补齐 → 按 (doc_id, chunk_index) 排序。
        任何异常抛 KBError（调用方捕获后降级，不扩散到出题流程）"""
        top_k = top_k or self._settings.kb_top_k
        min_score = min_score if min_score is not None else self._settings.kb_min_score
        # Chroma where 顶层只允许一个操作符：多条件必须用 $and 组合，双键 dict 会校验失败
        filt: dict = {"user_id": user_id}
        if kb_id is not None:
            filt = {"$and": [{"user_id": user_id}, {"kb_id": kb_id}]}
        try:
            vector = await asyncio.to_thread(self.embeddings.embed_query, query)
            res = await asyncio.to_thread(
                self.collection.query,
                query_embeddings=[vector],
                n_results=top_k * 2,
                where=filt,
                include=["documents", "metadatas", "distances"],
            )
        except Exception as exc:
            _logger.warning("%skb search fail user_id=%s kb_id=%s error=%s", task_id_kv(), user_id, kb_id, exc, exc_info=exc)
            raise KBError(f"知识库检索失败：{exc}") from exc
        ids = res.get("ids") or [[]]
        if not ids or not ids[0]:
            return []
        hits: list[ChunkHit] = []
        for i, chunk_id in enumerate(ids[0]):
            meta = res["metadatas"][0][i]
            score = 1.0 - res["distances"][0][i]  # cosine 空间：relevance = 1 - distance
            if score < min_score:
                continue
            hits.append(
                ChunkHit(
                    doc_id=meta["doc_id"],
                    kb_id=meta["kb_id"],
                    filename=meta.get("filename", ""),
                    chunk_index=meta["chunk_index"],
                    text=res["documents"][0][i],
                    score=score,
                )
            )
        return await self._expand_neighbors(hits)

    async def _expand_neighbors(self, hits: list[ChunkHit], radius: int = 1) -> list[ChunkHit]:
        """补齐命中块相邻 chunk。

        检索命中往往只有块内一小段，相邻块携带上下文保证资料语义完整；相邻块文本按
        确定性 id 二次 get 获取，缺失（文档边界）则跳过。
        """
        if not hits:
            return []
        needed: dict[tuple[int, int], ChunkHit | None] = {}
        for hit in hits:
            needed.setdefault((hit.doc_id, hit.chunk_index), hit)
        for hit in hits:
            for offset in range(1, radius + 1):
                for sign in (-1, 1):
                    needed.setdefault((hit.doc_id, hit.chunk_index + sign * offset), None)
        missing_ids = [
            f"doc_{doc_id}_{idx}" for (doc_id, idx), hit in needed.items() if hit is None
        ]
        if missing_ids:
            got = await asyncio.to_thread(self.collection.get, ids=missing_ids, include=["documents", "metadatas"])
            for i, cid in enumerate(got["ids"]):
                meta = got["metadatas"][i]
                needed[(meta["doc_id"], meta["chunk_index"])] = ChunkHit(
                    doc_id=meta["doc_id"],
                    kb_id=meta["kb_id"],
                    filename=meta.get("filename", ""),
                    chunk_index=meta["chunk_index"],
                    text=got["documents"][i],
                    score=0.0,
                )
        result = [hit for hit in needed.values() if hit is not None]
        result.sort(key=lambda h: (h.doc_id, h.chunk_index))  # 按文档内顺序（资料拼接语义连续）
        return result


def format_chunks_for_prompt(hits: list[ChunkHit], limit: int = 4000) -> str:
    """检索片段拼接为【文档资料】段落。片段超限截断，与联网资料同一上下文预算。"""
    parts = [
        f"【来源：{hit.filename} 第 {hit.chunk_index + 1} 段】\n{hit.text}" for hit in hits
    ]
    joined = "\n\n".join(parts)
    return joined[:limit]
